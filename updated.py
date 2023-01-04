import csv
from operator import index, sub
from matplotlib import pyplot as plt
import numpy as np
import docx 
from docx import Document 
from docx.shared import Inches
import os

main = open('sample.csv','r+')
contant = csv.reader(main)
obj = list(contant)
examName = (str(obj[0][0]))
subjects= list()
subject1 = list()
subject2 = list()
subject3 = list()
subject4 = list()
subject5 = list()
subject6 = list()
sub1pass = list()
sub2pass = list()
sub3pass = list()
sub4pass = list()
sub5pass = list()
sub6pass = list()
sub1fail = list()
sub2fail = list()
sub3fail = list()
sub4fail = list()
sub5fail = list()
sub6fail = list()

for i in range(3,len(obj)):
    for j in range(len(obj[3])):
        if j == 4:
            subject1.append(int(obj[i][j]))
        if j == 5:
            subject2.append(int(obj[i][j]))
        if j == 6:
            subject3.append(int(obj[i][j]))
        if j == 7:
            subject4.append(int(obj[i][j]))
        if j == 8:
            subject5.append(int(obj[i][j]))
        if j == 9:
            subject6.append(int(obj[i][j]))

for i in range(1,2):
    for j in range(len(obj[1])):
        if j == 4:
            subjects.append(obj[i][j])
        if j == 5:
            subjects.append(obj[i][j])
        if j == 6:
            subjects.append(obj[i][j])
        if j == 7:
            subjects.append(obj[i][j])
        if j == 8:
            subjects.append(obj[i][j])
        if j == 9:
            subjects.append(obj[i][j])
studentsNAME = list()
for i in range(3,len(obj)):
    for j in range(len(obj[3])):
        if j == 2:
            studentsNAME.append(obj[i][j])

    #if len(FAIL) > len(PASS):
def Average(subject):
    value = sum(subject)//len(subject)
    return value

def Percentage(value):
    perValue = value/40 * 100
    return perValue

def passCheck(subject,passVAl,failVAL,name):
    for i in range(len(subject)):
        if Percentage(subject[i]) >= 35 :
            passVAl.append(name[i])
        else :
            failVAL.append(name[i])
allimgPath = list()
def imgPathDEF(count):
    locationOFfile = os.getcwd()
    endOfpath = '/matplotlibExample.png'
    pathChanger =list(endOfpath)
    for i in range(len(pathChanger)):
        if pathChanger[i] == '.':
            b = pathChanger[i-1] + str(count)
            pathChanger[i-1] = b
    newPath = str() 
    for i in range(len(pathChanger)):
        newPath += pathChanger[i]
        allimgPath.append(locationOFfile + newPath)
    return locationOFfile + newPath



imgSAVEcount = list()
def genaratePIE(subject):
  
    pervalue = Percentage(subject)
    till35 = 0
    till50 = 0
    till75 = 0
    till90 = 0
    till95 = 0
    sentem = 0

    for i in range(len(pervalue)):
        if pervalue[i] in range(0,35 + 1):
            till35 += 1
        elif pervalue[i] in range(35,50+1):
            till50 += 1
        elif pervalue[i] in range(50,75+1):
            till75 += 1 
        elif pervalue[i] in range(76,90+1):
            till90 += 1
        elif pervalue[i] in range(91,100):
            till95 += 1
        elif pervalue[i] == 100:
            sentem += 1

    data = [till35,till50,till75,till90,till95,sentem]
    value = ['0-35%','36-50%','51-75%','76-90%','91-99%','100%']
    fig = plt.figure(figsize=(10,7))
    plt.pie(data,labels= value)
    plt.savefig(imgPathDEF(len(imgSAVEcount)))
    imgSAVEcount.append(1)
   # plt.show()
passCheck(subject1,sub1pass,sub1fail,studentsNAME)
passCheck(subject2,sub2pass,sub2fail,studentsNAME)
passCheck(subject3,sub3pass,sub3fail,studentsNAME)
passCheck(subject4,sub4pass,sub4fail,studentsNAME)
passCheck(subject5,sub5pass,sub5fail,studentsNAME)
passCheck(subject6,sub6pass,sub6fail,studentsNAME)
document = Document()

document.add_heading('Result Analysis', 0)
document.add_paragraph('Result Analysis of {}'.format(examName))
document.add_heading(subjects[0], level=1)
document.add_picture(allimgPath[0], width=Inches(5.25))
document.add_paragraph()
document.add_page_break()

document.save('demo.docx')

